from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Inches
from io import BytesIO
import json
import base64
import re

app = FastAPI()

def is_base64_image(data: str):
    return re.match(r'^data:image/(png|jpeg|jpg|gif|bmp|tiff);base64,', data)

@app.post("/generate-docx/")
async def generate_docx(
    template_file: UploadFile = File(...),
    json_data: str = Form(...)
):
    # テンプレートDOCXファイルを読み込み
    template_bytes = await template_file.read()
    doc = Document(BytesIO(template_bytes))

    # JSONデータをPythonの辞書に変換
    data = json.loads(json_data)

    # 段落ごとにプレースホルダーを処理
    for paragraph in doc.paragraphs:
        text = paragraph.text
        modified_text = text  # オリジナルのテキストを保持

        for key, value in data.items():
            placeholder = f"{{{{ {key} }}}}"

            if placeholder in text:
                # Base64エンコードされた画像の場合
                if isinstance(value, str) and is_base64_image(value):
                    image_data = base64.b64decode(value.split(",")[1])
                    modified_text = modified_text.replace(placeholder, "")  # プレースホルダーをクリア
                    paragraph.text = modified_text  # 先にテキストを置換
                    paragraph.add_run().add_break()  # 画像前の改行を追加
                    run = paragraph.add_run()
                    run.add_picture(BytesIO(image_data), width=Inches(4))
                    paragraph.add_run().add_break()  # 画像後の改行を追加

                # 値がリストの場合は、表として処理
                elif isinstance(value, list) and all(isinstance(i, list) for i in value):
                    modified_text = modified_text.replace(placeholder, "")
                    paragraph.text = modified_text  # 先にテキストを置換
                    paragraph.add_run().add_break()  # 表前の改行を追加
                    table = doc.add_table(rows=1, cols=len(value[0]))
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(value[0]):
                        hdr_cells[i].text = header
                    for row_data in value[1:]:
                        row_cells = table.add_row().cells
                        for j, cell_data in enumerate(row_data):
                            row_cells[j].text = cell_data
                    paragraph = doc.add_paragraph()  # 表後に新しい段落を追加

                # それ以外の場合は、単純なテキスト置換
                else:
                    modified_text = modified_text.replace(placeholder, str(value))

        # 最後に段落を更新
        paragraph.text = modified_text

    # DOCXファイルをメモリに保存
    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)

    # バイナリ形式でDOCXファイルを返す
    return StreamingResponse(output_buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=generated.docx"})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
